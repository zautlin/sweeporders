#!/usr/bin/env python3
"""
Generate comprehensive Word document from sweep order analysis markdown files.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path

def add_heading_with_style(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph_with_style(doc, text, bold=False, italic=False, color=None):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return para

def parse_markdown_table(lines):
    """Parse a markdown table into rows and columns"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|---'):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    return rows

def add_table_from_markdown(doc, table_lines):
    """Add a table to the document from markdown format"""
    rows = parse_markdown_table(table_lines)
    if len(rows) < 2:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, cell in enumerate(rows[0]):
        table.rows[0].cells[i].text = cell
        # Make header bold
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for row_idx in range(1, len(rows)):
        for col_idx, cell_text in enumerate(rows[row_idx]):
            table.rows[row_idx].cells[col_idx].text = cell_text
    
    return table

def process_markdown_content(doc, content, title):
    """Process markdown content and add to document"""
    lines = content.split('\n')
    i = 0
    
    # Add document title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip first heading (already added as title)
        if i < 5 and line.startswith('# '):
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'No Spacing'
                para.paragraph_format.left_indent = Inches(0.5)
                for run in para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            add_table_from_markdown(doc, table_lines)
            doc.add_paragraph()  # Add spacing
            table_lines = []
            in_table = False
        
        # Handle headings
        if line.startswith('## '):
            add_heading_with_style(doc, line[3:], level=1)
        elif line.startswith('### '):
            add_heading_with_style(doc, line[4:], level=2)
        elif line.startswith('#### '):
            add_heading_with_style(doc, line[5:], level=3)
        elif line.startswith('##### '):
            add_heading_with_style(doc, line[6:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()
        
        # Handle bold/italic text
        elif line.strip() and not line.startswith('#'):
            # Check for special formatting
            bold = '**' in line
            
            # Clean markdown formatting
            clean_line = line.replace('**', '').replace('*', '')
            
            # Handle special markers
            if line.startswith('- ✅') or line.startswith('- ❌') or line.startswith('- ⚠'):
                para = doc.add_paragraph(clean_line, style='List Bullet')
            elif line.startswith('- '):
                para = doc.add_paragraph(clean_line[2:], style='List Bullet')
            elif line.strip().startswith('|'):
                # Skip, handled by table logic
                pass
            else:
                para = doc.add_paragraph(clean_line)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)
    
    doc.add_page_break()

def main():
    """Main function to generate Word document"""
    print("Generating comprehensive Word document from markdown files...")
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Centre Point Sweep Order Analysis - Comprehensive Findings"
    doc.core_properties.author = "Quantitative Research Team"
    doc.core_properties.comments = "Analysis of 26,651 sweep orders across 4 securities"
    
    # Add cover page
    title = doc.add_heading("Centre Point Sweep Order Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Comprehensive Findings Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Analysis Date: January 2026\n")
    run.font.size = Pt(12)
    run = info.add_run("Dataset: 26,651 sweep orders across 4 securities\n")
    run.font.size = Pt(12)
    run = info.add_run("Period: September 5, 2024 (BHP, DRR, WTC) + May 5, 2024 (CBA)")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Key finding box
    key_finding = doc.add_paragraph()
    key_finding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    key_finding_run = key_finding.add_run(
        "KEY FINDING\n\n"
        "Current lit market resting strategy costs\n"
        "$3.4M-$3.6M annually\n"
        "compared to dark pool alternative"
    )
    key_finding_run.font.size = Pt(16)
    key_finding_run.font.bold = True
    key_finding_run.font.color.rgb = RGBColor(192, 0, 0)
    
    doc.add_page_break()
    
    # Read and process each markdown file
    docs_dir = Path('/Users/agautam/workspace/python/sweeporders/docs')
    
    files = [
        ('README.md', 'Overview and Navigation Hub'),
        ('EXECUTIVE_SUMMARY.md', 'Executive Summary'),
        ('STATISTICAL_SIGNIFICANCE.md', 'Statistical Significance Analysis'),
        ('METRICS_METHODOLOGY.md', 'Metrics Methodology'),
    ]
    
    for filename, section_title in files:
        filepath = docs_dir / filename
        print(f"Processing {filename}...")
        
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            process_markdown_content(doc, content, section_title)
            print(f"  ✓ {filename} processed successfully")
        except Exception as e:
            print(f"  ✗ Error processing {filename}: {e}")
    
    # Save document
    output_path = docs_dir / 'Sweep_Order_Analysis_Comprehensive_Report.docx'
    doc.save(output_path)
    
    print(f"\n✅ Document generated successfully!")
    print(f"📄 Saved to: {output_path}")
    print(f"📊 File size: {output_path.stat().st_size / 1024:.1f} KB")

if __name__ == '__main__':
    main()
